#!/usr/bin/env python3
"""Generate the workshop paper draft as a Word .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ============================================================
#                         TITLE
# ============================================================

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'Sparse Autoencoders Reveal How Training Objectives Shape\n'
    'Structural Representations in Protein Language Models'
)
run.bold = True
run.font.size = Pt(14)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Anonymous Authors')
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()  # spacer

# ============================================================
#                        ABSTRACT
# ============================================================

heading('Abstract', level=1)

doc.add_paragraph(
    'Protein language models (PLMs) trained with different objectives -- '
    'masked language modeling (ESM-2), autoregressive next-token prediction (ProtGPT2), '
    'and encoder-decoder denoising (ProtT5) -- achieve strong performance on downstream tasks, '
    'yet how their training objectives shape internal representations remains poorly understood. '
    'We apply TopK sparse autoencoders (SAEs) to four PLM architectures at five matched relative '
    'depths, training 20 SAEs on 1,500 SCOPe proteins with rigorous protein-level holdout validation. '
    'We find a stark dissociation: ESM-2 (bidirectional) learns features that are structurally local '
    '(Cohen\'s d = 0.6-1.9 above ProtGPT2 at every depth), while ProtGPT2 (causal) learns features '
    'that are sequentially local (d = 1.7-5.9 above ESM-2). This dissociation is robust across '
    'three random seeds (d std < 0.05) and two sparsity settings (k in {128, 256}). '
    'Within ProtT5, we discover a depth-dependent reversal: early decoder layers inherit the '
    'encoder\'s structural locality via cross-attention, but late decoder layers collapse to '
    'sequence-local prediction heads. These findings demonstrate that dictionary learning provides '
    'a principled lens for comparing how architectural choices imprint on learned protein representations.'
)

# ============================================================
#                      1. INTRODUCTION
# ============================================================

heading('1  Introduction', level=1)

doc.add_paragraph(
    'Protein language models (PLMs) have become foundational tools in computational biology, '
    'achieving state-of-the-art performance on tasks ranging from structure prediction to '
    'function annotation (Lin et al., 2023; Elnaggar et al., 2021; Ferruz et al., 2022). '
    'These models are trained with fundamentally different objectives: ESM-2 uses bidirectional '
    'masked language modeling (MLM), ProtGPT2 uses unidirectional next-token prediction, and '
    'ProtT5 uses an encoder-decoder denoising objective. Despite their shared domain (protein '
    'sequences), the training objective imposes distinct inductive biases on the learned '
    'representations -- biases that remain largely opaque.'
)

doc.add_paragraph(
    'Sparse autoencoders (SAEs) have emerged as a powerful tool for mechanistic interpretability '
    'in natural language processing (Bricken et al., 2023; Templeton et al., 2024; Gao et al., 2024), '
    'decomposing dense neural activations into sparse, interpretable features. Recent work has '
    'applied SAEs to protein models (Simon & Zou, 2024), but existing studies focus on single '
    'architectures (typically ESM-2) and emphasize feature discovery rather than systematic '
    'cross-architecture comparison.'
)

doc.add_paragraph(
    'We ask: how do different PLM training objectives shape the geometry of learned features, '
    'and can SAE-based dictionary learning reveal these differences? We train TopK SAEs (Gao et al., '
    '2024) on four PLM variants -- ESM-2, ProtGPT2, and both the encoder and decoder of ProtT5 -- '
    'at five matched relative depths (0%, 25%, 50%, 75%, 100% of total transformer blocks). '
    'We test five pre-registered hypotheses about structural locality, sequential locality, '
    'feature interpretability, encoder-decoder asymmetry, and depth trends.'
)

doc.add_paragraph(
    'Our contributions are: (1) the first systematic SAE-based comparison across PLM architectures '
    'at matched depths; (2) a clean empirical dissociation between structural and sequential feature '
    'locality driven by training objective (Cohen\'s d up to 5.9); (3) the discovery of a '
    'depth-dependent reversal in ProtT5 encoder vs. decoder feature properties; and (4) rigorous '
    'validation including protein-level holdout, multi-seed reproducibility, and hyperparameter '
    'robustness checks.'
)

# ============================================================
#                      2. METHODS
# ============================================================

heading('2  Methods', level=1)

# 2.1 Dataset
heading('2.1  Dataset', level=2)

doc.add_paragraph(
    'We curate 1,500 proteins from SCOPe 2.08 (Chandonia et al., 2022) at 40% sequence identity '
    'filtering, stratified by fold class to preserve structural diversity (432 folds across all 7 '
    'SCOPe classes). We filter to proteins with >= 80% DSSP secondary-structure coverage and '
    'available PDB structures. Protein lengths range from 50 to 794 residues (mean 197, total '
    '295,240 residues). We apply a deterministic protein-level 90/10 train/validation split '
    '(seed 42, 1,350 train / 150 validation proteins), fixed across all models and seeds '
    'to ensure fair cross-architecture comparison of holdout generalization.'
)

# 2.2 PLM Architectures
heading('2.2  PLM Architectures and Depth Matching', level=2)

doc.add_paragraph(
    'We compare four PLM variants spanning three architectural families:'
)

add_table(
    ['Model', 'Architecture', 'Params', 'Objective', 'Layers probed'],
    [
        ['ESM-2 (t33)', 'Bidirectional encoder', '650M', 'Masked LM', '0, 8, 16, 24, 32'],
        ['ProtGPT2', 'Causal decoder', '738M', 'Next-token', '0, 9, 18, 27, 35'],
        ['ProtT5-enc', 'Bidir. encoder (seq2seq)', '~1.2B', 'Denoising enc.', '0, 6, 12, 18, 23'],
        ['ProtT5-dec', 'Autoreg. decoder (seq2seq)', '~3B*', 'Denoising dec.', '0, 6, 12, 18, 23'],
    ]
)

doc.add_paragraph(
    '*ProtT5 decoder requires the full encoder-decoder model (~3B total). '
    'Layers are selected at 0%, 25%, 50%, 75%, and 100% relative depth to enable '
    'fair cross-architecture comparison despite different total layer counts.'
)

doc.add_paragraph(
    'Importantly, the ProtT5 decoder is not strictly causal: each decoder position cross-attends '
    'to the full bidirectional encoder output, giving it access to "future" residue information. '
    'We therefore frame the ESM-2 vs. ProtGPT2 comparison as the pure bidirectional/causal test '
    '(H1, H2), and the ProtT5 encoder vs. decoder comparison as testing unconstrained encoding '
    'vs. autoregressive decoding with global context (H4).'
)

# 2.3 SAE Training
heading('2.3  TopK Sparse Autoencoder Training', level=2)

doc.add_paragraph(
    'We train TopK SAEs following the methodology of Gao et al. (2024). Each SAE has expansion '
    'factor 8 (hidden dimension = 8 x input dimension), yielding 10,240 features for ESM-2/ProtGPT2 '
    '(input dim 1,280) and 8,192 features for ProtT5 (input dim 1,024). We set k_sparse = 256, '
    'selected via an ablation on ESM-2 layer 16 as the value achieving the smallest train/validation '
    'explained-variance gap (0.092) among k in {64, 128, 256} (Appendix A).'
)

doc.add_paragraph(
    'Input normalization. PLM hidden states exhibit outlier features -- a small number of dimensions '
    'with activations 50-100x the typical scale (cf. Dettmers et al., 2022). Without normalization, '
    'these outliers dominate the MSE loss and destabilize TopK SAE training, producing negative '
    'explained variance and loss U-curves. Following Bricken et al. (2023), we rescale all '
    'embeddings so that the mean per-token L2 norm equals sqrt(D), computed on the training set '
    'and applied identically to the validation set.'
)

doc.add_paragraph(
    'Training details. We train for 60 epochs with AdamW (lr = 5e-5, cosine decay to 5e-6), '
    'batch size 4,096, auxiliary loss coefficient 1/32 with k_aux = 64, and dead-latent threshold '
    '1,000,000 tokens. Decoder columns are normalized to unit L2 norm after each optimizer step. '
    'The pre-encoder bias b_pre is set to the training-set mean. We report explained variance (EV) '
    'on both training and held-out validation proteins; all 20 SAEs achieve positive validation EV '
    '(range 0.66-0.99) with train-val gaps below 0.10.'
)

doc.add_paragraph(
    'Reproducibility. We repeat the full pipeline at three random seeds (42, 43, 44), varying '
    'only the SAE weight initialization and dataloader shuffle order while holding the protein-level '
    'split fixed. Cross-seed standard deviation of Cohen\'s d for H1/H2 is below 0.05 at every '
    'depth (Table 2), confirming that the findings are not sensitive to SAE initialization.'
)

# 2.4 Structural / Sequential Locality
heading('2.4  Structural and Sequential Locality', level=2)

doc.add_paragraph(
    'For each SAE feature, we compute a structural locality score (struct_delta) and a sequential '
    'locality score (seq_delta), following a shuffle-controlled Cohen\'s d framework:'
)

doc.add_paragraph(
    '(1) Structural neighbors are defined as residue pairs with C-alpha distance < 8 Angstrom and '
    'sequence separation >= 12 residues (long-range contacts, following standard conventions in '
    'protein contact prediction). (2) Sequential neighbors are residues within +/-2 positions in '
    'sequence. (3) For each feature, we compute the mean activation of structural (or sequential) '
    'neighbors at "active" residues (top 10% of activation values). (4) We compare this to a '
    'shuffled baseline (5 within-protein permutations) via Cohen\'s d. (5) The final delta = '
    'observed Cohen\'s d - shuffled Cohen\'s d, capturing locality above chance.'
)

doc.add_paragraph(
    'Positive struct_delta means the feature fires preferentially on residues whose 3D neighbors '
    'also express the feature (structural locality). Positive seq_delta means the feature fires '
    'on residues whose sequence neighbors also express the feature (sequential locality). '
    'We compute these via sparse matrix multiplication for efficiency.'
)

# 2.5 Feature Interpretability
heading('2.5  Feature Interpretability', level=2)

doc.add_paragraph(
    'For each feature, we compute Pearson correlations with three structural labels: helix '
    '(DSSP H/G/I), strand (E/B), and burial (neighbor count at 8 Angstrom). We apply '
    'Benjamini-Hochberg FDR correction per label and define a feature as "interpretable" if '
    'any of its three q-values falls below a threshold. We report results at multiple thresholds '
    '(q < 0.05 through q < 1e-6) to characterize the full sensitivity profile.'
)

# ============================================================
#                     3. HYPOTHESES
# ============================================================

heading('3  Hypotheses', level=1)

hypotheses = [
    ('H1', 'ESM-2 features have higher structural locality (struct_delta) than ProtGPT2 features '
           'at every matched relative depth.'),
    ('H2', 'ProtGPT2 features have higher sequential locality (seq_delta) than ESM-2 features '
           'at every matched relative depth.'),
    ('H3', 'ESM-2 features are more interpretable (higher fraction significantly correlated with '
           'secondary structure or burial) than ProtGPT2 features.'),
    ('H4', 'ProtT5 encoder features have higher structural locality (H4a) and interpretability '
           '(H4c) than ProtT5 decoder features, while ProtT5 decoder features have higher '
           'sequential locality (H4b). [Framed as unconstrained encoding vs. autoregressive '
           'decoding, not pure bidirectional vs. causal.]'),
    ('H5', 'Structural locality increases with depth across all models.'),
]

for tag, text in hypotheses:
    p = doc.add_paragraph()
    run = p.add_run(f'{tag}: ')
    run.bold = True
    p.add_run(text)

# ============================================================
#                       4. RESULTS
# ============================================================

heading('4  Results', level=1)

# 4.1 H1/H2
heading('4.1  Bidirectional vs. Causal: A Clean Dissociation (H1, H2)', level=2)

doc.add_paragraph(
    'H1 and H2 are both supported at all five matched depths with large effect sizes (Table 1). '
    'ESM-2 features exhibit positive structural locality (struct_delta > 0 for 60-88% of features), '
    'while ProtGPT2 features are structurally anti-local (struct_delta < 0 for 66-92% of features). '
    'Conversely, ProtGPT2\'s sequential locality is 5-10x larger than ESM-2\'s across all depths. '
    'The effect sizes are enormous: Cohen\'s d for H2 reaches +5.9 at the final layer, indicating '
    'that the causal training objective imprints a strong sequential-locality signature that persists '
    'even in late-layer representations.'
)

para('Table 1: H1 and H2 effect sizes at matched relative depths (mean +/- std across 3 seeds).', bold=True, size=9)

add_table(
    ['Depth', 'H1 Cohen\'s d', 'H1 sig?', 'H2 Cohen\'s d', 'H2 sig?'],
    [
        ['0%',   '+1.91 +/- 0.01', '5/5', '+3.24 +/- 0.01', '5/5'],
        ['25%',  '+1.24 +/- 0.02', '5/5', '+2.55 +/- 0.03', '5/5'],
        ['50%',  '+1.39 +/- 0.01', '5/5', '+1.69 +/- 0.00', '5/5'],
        ['75%',  '+0.62 +/- 0.01', '5/5', '+2.37 +/- 0.01', '5/5'],
        ['100%', '+1.72 +/- 0.04', '5/5', '+5.91 +/- 0.01', '5/5'],
    ]
)

doc.add_paragraph(
    'Robustness. Both H1 and H2 are confirmed under an alternative sparsity setting (k_sparse = 128): '
    'all 10 contrasts remain significant with preserved effect direction. Cohen\'s d values are '
    'somewhat smaller at k = 128 (H1 range [0.42, 1.55], H2 range [1.80, 4.92]) but remain in '
    'the "large effect" regime at every depth.'
)

# 4.2 H3
heading('4.2  Feature Interpretability (H3)', level=2)

doc.add_paragraph(
    'At the default threshold (q < 0.05), ESM-2 has higher interpretability than ProtGPT2 at all '
    '5 depths (95-99% vs. 92-98%), but both models approach a ceiling. To escape this ceiling, we '
    'evaluate at progressively tighter thresholds (Table 2). At q < 1e-6, the gap widens to '
    '12-27 percentage points, with the largest gap at the 50% depth (ESM-2: 88.7% vs. ProtGPT2: '
    '61.7%, chi-squared p ~ 0). At k_sparse = 128, H3 is supported at 4/5 depths, with a marginal '
    'reversal at the input layer (0% depth).'
)

para('Table 2: H3 interpretability gap (ESM-2 % - ProtGPT2 %) at varying q-thresholds.', bold=True, size=9)

add_table(
    ['Depth', 'q < 0.05', 'q < 0.001', 'q < 1e-6'],
    [
        ['0%',   '+2.7 pp', '+7.0 pp',  '+12.0 pp'],
        ['25%',  '+5.0 pp', '+13.3 pp', '+21.8 pp'],
        ['50%',  '+4.9 pp', '+15.0 pp', '+27.0 pp'],
        ['75%',  '+1.4 pp', '+5.5 pp',  '+11.1 pp'],
        ['100%', '+3.7 pp', '+12.1 pp', '+24.8 pp'],
    ]
)

# 4.3 H4
heading('4.3  ProtT5 Encoder vs. Decoder: A Depth-Dependent Reversal (H4)', level=2)

doc.add_paragraph(
    'H4 reveals a nuanced, layer-dependent pattern rather than a clean encoder-decoder split. '
    'At early layers (0%, 25%), the decoder outperforms the encoder on all three sub-hypotheses: '
    'it has higher structural locality (H4a: d = -0.30 to -0.38, decoder wins), higher sequential '
    'locality (H4b: d = +0.48 to +1.64, decoder wins), and higher interpretability (H4c: 98% vs. '
    '87% at layer 0). By the mid-to-late layers (50-100%), this pattern reverses: the encoder '
    'takes over structural dominance (H4a: d = +0.40 to +0.50), while the decoder\'s sequential '
    'locality advantage disappears or inverts.'
)

doc.add_paragraph(
    'We interpret this as follows: at early layers, the decoder inherits the encoder\'s '
    'bidirectional context via cross-attention, producing representations that are at least as '
    'structurally rich as the encoder\'s own. At late layers, the autoregressive decoding objective '
    'dominates, collapsing decoder representations toward sequence-local prediction heads while '
    'the encoder maintains its structural encoding. This transition is consistent with the known '
    'behavior of encoder-decoder models, where cross-attention influence diminishes relative to '
    'self-attention in deeper layers.'
)

# 4.4 H5
heading('4.4  Depth Trends in Structural Locality (H5)', level=2)

doc.add_paragraph(
    'H5 is model-dependent. Using per-feature Spearman correlation between layer index and '
    'struct_delta across all features and layers (N = 40,960-51,200 per model):'
)

add_table(
    ['Model', 'Spearman rho', 'p-value', 'Direction'],
    [
        ['ESM-2',      '-0.233', '< 1e-300', 'Decreases (contradicts H5)'],
        ['ProtGPT2',   '+0.272', '< 1e-300', 'Increases (supports H5)'],
        ['ProtT5-enc', '+0.371', '< 1e-300', 'Increases (supports H5)'],
        ['ProtT5-dec', '-0.054', '2.6e-28',  'Flat / slightly decreasing'],
    ]
)

doc.add_paragraph(
    'ESM-2\'s structural locality is front-loaded: it peaks at layer 0 (struct_delta = 0.143) '
    'and drops sharply in deeper layers (0.019-0.053). This is consistent with the MLM objective, '
    'which can exploit structural context immediately via bidirectional attention. In contrast, '
    'causal models (ProtGPT2, ProtT5-enc) must build up structural representations incrementally '
    'through layers, producing the positive depth trend that H5 predicts. This model-dependent '
    'asymmetry is itself a finding: the training objective determines not just the magnitude '
    'of structural encoding but its depth profile.'
)

# 4.5 Generalization
heading('4.5  Generalization and Reproducibility', level=2)

doc.add_paragraph(
    'All 20 SAEs achieve positive validation explained variance (range 0.66-0.99) on 150 '
    'held-out proteins, with train-val EV gaps below 0.10 at k = 256 (range 0.000-0.092). '
    'At k = 128, ESM-2 mid-layer gaps slightly exceed 0.10 (0.106-0.139), consistent with the '
    'ablation finding that k = 256 provides better generalization for ESM-2\'s high-effective-rank '
    'representations. Across three random seeds, Cohen\'s d standard deviation is below 0.05 for '
    'all H1/H2 contrasts, confirming that the cross-architecture dissociation is not an artifact '
    'of SAE initialization.'
)

# 4.6 Causal Evidence (TBD)
heading('4.6  Causal Evidence via Activation Clamping (Preliminary)', level=2)

doc.add_paragraph(
    '[This section will report results from activation clamping experiments on ESM-2 layer 16, '
    'where the top-10 structural features identified by struct_delta are ablated or amplified '
    'during the ESM-2 forward pass, and the effect on zero-shot contact prediction precision '
    'is measured. Experiments are in progress and will be included in the final version.]'
)

# ============================================================
#                     5. DISCUSSION
# ============================================================

heading('5  Discussion', level=1)

doc.add_paragraph(
    'Our results establish that sparse dictionary learning can reveal systematic differences '
    'in how PLM training objectives shape internal representations. The H1/H2 dissociation -- '
    'bidirectional models encode structural locality, causal models encode sequential locality -- '
    'is clean, large, and robust. This finding has practical implications: if downstream tasks '
    'require structural reasoning (e.g., contact prediction, fold classification), bidirectional '
    'models may provide more directly useful intermediate features, while causal models may '
    'excel at tasks requiring local sequence context (e.g., generation, perplexity-based scoring).'
)

doc.add_paragraph(
    'The ProtT5 encoder-decoder reversal (H4) adds nuance: the decoder\'s cross-attention to the '
    'encoder provides access to global structural context, but the autoregressive training objective '
    'progressively overrides this signal in deeper layers. This depth-dependent transition has not, '
    'to our knowledge, been previously documented in protein language models.'
)

doc.add_paragraph(
    'The failure of H5 for ESM-2 (structural locality decreasing with depth) is itself informative. '
    'It suggests that ESM-2\'s bidirectional attention enables structural feature formation at the '
    'earliest layers, after which deeper layers may shift toward more abstract or fold-level '
    'representations that are less locally structural. This contrasts with the progressive '
    'structure-building observed in causal models and is consistent with the "information '
    'bottleneck" view of transformer depth.'
)

# Limitations
heading('5.1  Limitations', level=2)

doc.add_paragraph(
    'Our study has several limitations. (1) We analyze 1,500 proteins, which is small relative '
    'to PLM pretraining corpora; however, our protein-level holdout and multi-seed analysis '
    'confirm generalization. (2) Our structural/sequential locality metrics are correlational; '
    'activation clamping experiments (Section 4.6) will provide causal evidence. (3) We compare '
    'models of different sizes (650M-3B parameters), which confounds architecture with scale. '
    '(4) We use a single expansion factor (8x) and do not explore the effect of SAE capacity. '
    '(5) ProtGPT2 uses BPE tokenization, requiring a token-to-residue projection that may '
    'introduce alignment noise.'
)

# ============================================================
#                    6. RELATED WORK
# ============================================================

heading('6  Related Work', level=1)

doc.add_paragraph(
    'Sparse autoencoders for interpretability. Bricken et al. (2023) introduced SAEs for '
    'mechanistic interpretability, decomposing MLP activations in a small transformer into '
    'interpretable features. Templeton et al. (2024) scaled this approach to Claude. Gao et al. '
    '(2024) introduced TopK SAEs with the auxiliary loss for dead-latent recovery that we adopt. '
    'Our work extends this methodology to the protein domain and, crucially, to cross-architecture '
    'comparison.'
)

doc.add_paragraph(
    'SAEs on protein models. Simon & Zou (2024, InterPLM) applied SAEs to ESM-2, discovering '
    'features for binding sites and post-translational modifications. Their work focuses on '
    'single-model feature discovery; we complement it with a systematic cross-architecture '
    'hypothesis-driven comparison.'
)

doc.add_paragraph(
    'Protein language models. ESM-2 (Lin et al., 2023), ProtGPT2 (Ferruz et al., 2022), and '
    'ProtT5 (Elnaggar et al., 2021) represent three major PLM families. Prior comparisons have '
    'focused on downstream task performance (e.g., Rao et al., 2019; Elnaggar et al., 2021) '
    'rather than interpretability of internal representations.'
)

doc.add_paragraph(
    'Outlier features in transformers. Dettmers et al. (2022) documented emergent outlier '
    'features in large language models that dominate hidden-state magnitudes. We encounter the '
    'same phenomenon in all four PLMs and address it via Bricken-style input normalization, '
    'which we find is necessary for stable SAE training on protein embeddings.'
)

# ============================================================
#                    7. CONCLUSION
# ============================================================

heading('7  Conclusion', level=1)

doc.add_paragraph(
    'We have shown that TopK sparse autoencoders, applied systematically across four protein '
    'language model architectures at matched relative depths, reveal a clear dissociation between '
    'bidirectional and causal training objectives: bidirectional models learn structurally local '
    'features while causal models learn sequentially local features. This finding is robust across '
    'random seeds, sparsity hyperparameters, and all five probed depth levels. Within the '
    'encoder-decoder ProtT5 architecture, we discover a depth-dependent reversal in which early '
    'decoder layers inherit structural locality from the encoder via cross-attention, while late '
    'layers collapse to sequence-local representations. These results demonstrate that dictionary '
    'learning is a powerful and principled tool for understanding how architectural and training '
    'choices shape the representations learned by domain-specific foundation models.'
)

# ============================================================
#                    REFERENCES
# ============================================================

heading('References', level=1)

refs = [
    'Bricken, T., et al. (2023). Towards Monosemanticity: Decomposing Language Models With Dictionary Learning. Anthropic.',
    'Chandonia, J.-M., et al. (2022). SCOPe: improvements to the structural classification of proteins -- extended database. NAR.',
    'Dettmers, T., et al. (2022). LLM.int8(): 8-bit Matrix Multiplication for Transformers at Scale. NeurIPS.',
    'Elnaggar, A., et al. (2021). ProtTrans: Toward Understanding the Language of Life Through Self-Supervised Learning. IEEE TPAMI.',
    'Ferruz, N., et al. (2022). ProtGPT2 is a deep unsupervised language model for protein design. Nature Communications.',
    'Gao, L., et al. (2024). Scaling and Evaluating Sparse Autoencoders. arXiv:2406.04093.',
    'Lin, Z., et al. (2023). Evolutionary-scale prediction of atomic-level protein structure with a language model. Science.',
    'Rao, R., et al. (2019). Evaluating Protein Transfer Learning with TAPE. NeurIPS.',
    'Simon, E. & Zou, J. (2024). InterPLM: Discovering Interpretable Features in Protein Language Models via Sparse Autoencoders. bioRxiv.',
    'Templeton, A., et al. (2024). Scaling Monosemanticity: Extracting Interpretable Features from Claude 3 Sonnet. Anthropic.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(9)

# ============================================================
#                    APPENDIX A
# ============================================================

doc.add_page_break()

heading('Appendix A: Sparsity Ablation', level=1)

doc.add_paragraph(
    'We select k_sparse via an ablation on ESM-2 layer 16 (50% relative depth), training SAEs '
    'at k in {64, 128, 256} with all other hyperparameters held fixed. We report explained '
    'variance on both training and held-out validation proteins:'
)

add_table(
    ['k_sparse', 'Train EV', 'Val EV', 'Gap'],
    [
        ['64',  '0.746', '0.612', '0.134'],
        ['128', '0.775', '0.658', '0.117'],
        ['256', '0.809', '0.717', '0.092'],
    ]
)

doc.add_paragraph(
    'k = 256 achieves the highest validation EV and the smallest train-val gap (0.092, the only '
    'value below 0.10). The counterintuitive finding that higher k improves generalization '
    'suggests that k = 64 and k = 128 are too constrained for ESM-2\'s effective intrinsic '
    'dimensionality, forcing the SAE to memorize specific training activations rather than learning '
    'generalizable features.'
)

# ============================================================
#                    APPENDIX B
# ============================================================

heading('Appendix B: Multi-Seed Reproducibility', level=1)

doc.add_paragraph(
    'We repeat the full pipeline at three random seeds (42, 43, 44), varying only the SAE weight '
    'initialization. Cross-seed standard deviations for all metrics are small:'
)

add_table(
    ['Metric', 'Max std across 20 layers'],
    [
        ['Train EV', '0.008'],
        ['Val EV', '0.002'],
        ['EV gap', '0.011'],
        ['H1 Cohen\'s d (per depth)', '0.044'],
        ['H2 Cohen\'s d (per depth)', '0.031'],
    ]
)

doc.add_paragraph(
    'The largest variability is in ESM-2 H1 at the 100% depth (d std = 0.044). All other '
    'contrasts have d std < 0.02. The cross-architecture dissociation is robust to initialization.'
)

# ============================================================
#                    SAVE
# ============================================================

out_path = 'paper_draft_mechinterp_workshop.docx'
doc.save(out_path)
print(f'Saved to {out_path}')

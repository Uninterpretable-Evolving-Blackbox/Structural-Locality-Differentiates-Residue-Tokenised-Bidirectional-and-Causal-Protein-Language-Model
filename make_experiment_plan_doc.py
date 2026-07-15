#!/usr/bin/env python3
"""
make_experiment_plan_doc.py — generate the full-paper experiment plan (.docx)
=============================================================================

Produces EXPERIMENT_PLAN.docx: every planned experiment for the arXiv full
paper, with scientific justification, and a robustness / seed-repeat protocol
foregrounded and attached to each experiment.

Run:
  .venv/bin/python make_experiment_plan_doc.py
"""

from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1A, 0x57, 0x9C)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ----- base styling -----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def para(text="", italic=False, bold=False, color=None, size=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def kv(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(value)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    return t


# =====================================================================
#  TITLE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SAE-PLM Full-Paper Experiment Plan")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Grounding structural-locality features in biology, causality, and robust statistics")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Extension plan for the arXiv full paper · drafted {date.today().isoformat()}").font.size = Pt(9)

doc.add_paragraph()

# =====================================================================
#  1. PURPOSE
# =====================================================================
heading("1. Purpose and scope", 1)
para(
    "The accepted workshop paper establishes one robust result: bidirectional "
    "residue-tokenised PLMs (ESM-2) show stronger structural locality (L_struct) "
    "than causal residue-tokenised PLMs (RITA) at every matched depth. Reviewers "
    "and our own assessment converge on a single deeper limitation: L_struct is a "
    "novel, homemade co-activation statistic that has not been shown to (a) align "
    "with curated biological concepts, (b) be necessary relative to simpler "
    "baselines, or (c) causally influence model behaviour. This document specifies "
    "the experiments that turn the workshop result into a full paper by grounding "
    "the claim in field-standard interpretability metrics, causal interventions, "
    "and a robustness-first statistical protocol.")
para(
    "Design philosophy: the central claim must rest on CONVERGENT EVIDENCE from "
    "several independent lenses (concept alignment, linear decodability, causal "
    "intervention) rather than on one metric, and every reported effect must "
    "survive the seed-repeat and robustness battery in Section 2 BEFORE it is "
    "believed.", italic=True)

# =====================================================================
#  2. ROBUSTNESS PROTOCOL  (foregrounded)
# =====================================================================
heading("2. Robustness and reproducibility protocol (applies to every experiment)", 1)
para(
    "This protocol is the backbone of the paper. No experiment below is considered "
    "complete until its headline effect has been run through the relevant subset of "
    "these checks. The thresholds and seeds reuse the conventions already validated "
    "in the workshop paper so the new experiments inherit the same rigour.")

heading("2.1 Seed repeats (primary robustness axis)", 2)
bullets([
    ("SAE initialisation seeds", "Every SAE is retrained at seeds {42, 43, 44}. We report the cross-seed mean and standard deviation of the headline statistic (concept-F1, probe AUROC, causal effect, L_struct). Acceptance: direction preserved at all three seeds and cross-seed SD small relative to the effect (workshop precedent: SD <= 0.044 on H1 d)."),
    ("Why three seeds", "TopK SAEs are near-deterministic in our setting, so 3 seeds is sufficient to bound initialisation variance without excessive compute; high-EV regimes (e.g. RITA L3) are additionally flagged for cross-seed verification before any feature-level claim."),
    ("Data-resampling seeds", "Protein-level cluster bootstrap (B = 1000) resamples proteins with replacement to give 95% percentile CIs on every cross-model effect size; the bootstrap mean is reported as the conservative point estimate."),
    ("Null seeds", "Degree-matched and global-shuffle graph nulls, and the randomized-weights control, are each run at >= 3 random seeds to bound the floor."),
])

heading("2.2 Standard robustness battery", 2)
table(
    ["Axis", "Settings", "Acceptance criterion"],
    [
        ["SAE seed", "{42, 43, 44}", "direction at all seeds; small cross-seed SD"],
        ["Sparsity k", "{128, 256}", "direction preserved"],
        ["Train/val split", "90/10 (seed 42); 99/1; fresh 1,500 SCOPe subsample", "direction preserved; |Δ| small vs main"],
        ["Metric sweep", "Cα {6,8,10} Å; sep {8,12,24}; quantile {5,10,20}%; window {±1,±2,±4}", "direction across the grid (report cell counts)"],
        ["Held-out confirmation", "150-protein val-only recompute", "effect preserved on unseen proteins"],
        ["Statistics", "protein-level cluster bootstrap B=1000; BH-FDR for many-feature tests", "95% CI excludes 0; FDR < 0.05"],
        ["Negative control", "randomized-PLM-weights SAE; shuffled graphs", "effect collapses to floor"],
    ])
para(
    "Multiple-comparison control: any analysis that scans thousands of features "
    "(concept-F1, interpretability correlations) reports Benjamini–Hochberg FDR, "
    "and concept selection uses a held-out val→test split so reported F1 is never "
    "computed on the data used to pick the feature/threshold.", after=4)

# =====================================================================
#  3. EXPERIMENT CATALOGUE
# =====================================================================
heading("3. Experiment catalogue", 1)
para(
    "Each experiment lists its question/hypothesis, method (with the SOTA source "
    "it follows), why it is scientifically reasonable, its seed-repeat and "
    "robustness plan, expected outputs, feasibility on the target hardware "
    "(MacBook M5 Max, 128 GB unified memory), and risks. Experiments tagged "
    "[BUILT] already have validated scripts in the repository.", italic=True)

experiments = [
    {
        "id": "E0",
        "title": "Concept-F1: align features to curated biology  [BUILT: experiment_concept_f1.py]",
        "question": "Do SAE features correspond to real biological concepts, and does ESM-2 align more strongly than RITA?",
        "hypothesis": "ESM-2 has more concept-aligned features (F1 > 0.5) and higher mean best-F1 per concept than RITA at matched depths.",
        "method": "InterPLM (Simon & Zou, Nat. Methods 2025) domain-aware metric: precision = TP/(TP+FP) per amino acid; recall = DomainsWithTruePositive/TotalDomains per domain; F1 = harmonic mean. Binarise each feature over thresholds {0,0.15,0.5,0.6,0.8}; select best feature+threshold per concept on a concept-VAL protein split; report F1 on a held-out concept-TEST split. Concepts: SCOPe class/fold/superfamily/family (from scope_40.fa), DSSP secondary structure, and RSA bins.",
        "rationale": "Domain-aware recall is the field-standard fix for the problem that features are often more specific than annotations; it directly answers the 'is L_struct a real feature?' critique by testing against external, curated labels rather than a self-defined statistic. Pilot (512/10240 features) already recovers fold detectors with test F1 up to 0.998.",
        "robustness": "Run at all 9 ESM-2 and 9 RITA depths; repeat at SAE seeds {42,43,44} and k=128; recompute on the fresh SCOPe subsample; report BH-FDR-controlled counts and val→test F1 (no selection leakage); negative control C1 (randomized weights) must give near-zero concept-F1.",
        "outputs": "Per-concept best-feature F1 table; #features with F1>0.5 (val&test) per model/depth; ESM-2 vs RITA concept-F1 gap vs depth; Spearman(concept-F1, L_struct) scatter (tests whether the two metrics agree).",
        "feasibility": "~8 min/layer on cached activations; ~2.5 h for all 18 layers. CPU-bound, no model loading. HIGH.",
        "risks": "Rare SCOPe families have few domains; mitigated by the min-domains>=10 filter and by aggregating at fold/superfamily level.",
    },
    {
        "id": "C1",
        "title": "Randomized-weights negative control  [PLANNED]",
        "question": "Do concept alignment and L_struct require a TRAINED PLM, or do they arise from architecture/data priors alone?",
        "hypothesis": "An SAE trained on a randomly-initialised PLM yields near-zero concept-F1 and floor-level L_struct.",
        "method": "InterPLM shuffled-weights control: initialise ESM-2 (and RITA) with random weights, extract embeddings, train an SAE with identical hyperparameters, then run E0 and E4 unchanged.",
        "rationale": "Establishes that the bidirectional-vs-causal dissociation is a property of LEARNED representations, not of tokenisation or random projection geometry. This is the single most important control for the whole paper.",
        "robustness": "Repeat at >= 3 weight-init seeds; compare against trained models at matched depth/k; report the full concept-F1 and L_struct distributions, not just means.",
        "outputs": "Side-by-side trained-vs-random concept-F1 counts and L_struct histograms; demonstration that both collapse to floor under randomisation.",
        "feasibility": "~1–2 h/layer (embedding extraction + SAE training on MPS). Needs a ~15-line randomized-model extractor; reuses train_sae. MEDIUM.",
        "risks": "Random-weight activations may be ill-scaled; mitigated by the existing Bricken normalisation already in the pipeline.",
    },
    {
        "id": "E1",
        "title": "SAE-vs-raw linear probe: are SAEs necessary?  [BUILT: experiment_probe_baseline.py]",
        "question": "Does a linear probe on raw activations reproduce the ESM-2>RITA structural gap, and do SAE features beat raw for structure decodability?",
        "hypothesis": "Both raw and SAE probes show ESM-2>RITA on structure; SAE features match or exceed raw (esp. at layer 0 where whole-vector cosine is blind).",
        "method": "Adams et al. (ICML 2025): logistic probes for helix/strand/burial (and optional long-range contact) on raw vs SAE features, protein-split train/test, reporting F1 and AUROC at matched depths for both models.",
        "rationale": "Directly answers reviewer U2Yp ('do you even need SAEs?'). Either outcome strengthens the paper: raw-probe agreement shows H1 is method-robust; SAE advantage shows SAE-specific value. Pilot already shows SAE >= raw on strand/burial AUROC.",
        "robustness": "All 9 matched depths x 2 models; SAE seeds {42,43,44}; probe re-fit at 3 train/test protein splits; report mean +/- SD AUROC; identical preprocessing (z-scored raw, fixed C, balanced classes).",
        "outputs": "AUROC/F1 vs depth, raw vs SAE, ESM-2 vs RITA; layer-0 panel highlighting the cosine-blind regime.",
        "feasibility": "Re-extracts + caches raw embeddings (~10–20 min/model on MPS); probe fits in minutes. HIGH.",
        "risks": "Probe capacity confound (SAE has more dims); mitigated by L2 regularisation, matched train size, and reporting AUROC (threshold-free).",
    },
    {
        "id": "E2",
        "title": "Causal feature identification + causal H1  [EXTENDS experiment_activation_clamping.py]",
        "question": "Which features causally affect a structure readout, and does the bidirectional-vs-causal dissociation hold on the causal subset?",
        "hypothesis": "Ablating high-L_struct features degrades contact prediction more than random features; ESM-2 has more causally-structural features than RITA.",
        "method": "Forward-pass hook at matched depth: encode hidden state through SAE, ablate (clamp to 0) or amplify target features, decode, propagate. Score each feature's causal effect as the drop in long-range top-L/5 contact precision; define the causal subset by significant drop; re-test H1 on that subset; characterise causal features with E0 concept-F1.",
        "rationale": "Moves the paper from correlation to causation (reviewer 6P8E's top request) and isolates the features that actually matter, addressing the user's priority that analysis should focus on causal features.",
        "robustness": "Run for both ESM-2 and RITA; SAE seeds {42,43,44}; paired per-protein bootstrap CI on the precision drop; random-feature control of matched count and activation density; upgrade the readout from APC-cosine to a supervised contact probe (from E1) for a defensible absolute metric; intervention run on the 150 held-out proteins.",
        "outputs": "Per-feature causal effect ranking; causal-subset H1 effect sizes vs depth; fraction of causally-structural features (ESM-2 vs RITA); concept-F1 of the causal subset.",
        "feasibility": "Forward passes over ~150–200 proteins x features x conditions; CPU recommended (known MPS fp16-hook issue). Overnight. MEDIUM.",
        "risks": "Downstream propagation noise; mitigated by paired (same-protein) baseline-vs-intervention comparison and bootstrap CIs.",
    },
    {
        "id": "E3",
        "title": "Steering / activation clamping on masked-token probabilities  [EXTENDS clamping infra]",
        "question": "Do high-L_struct / high-concept-F1 features causally steer model outputs in interpretable ways?",
        "hypothesis": "Clamping a structural feature on one residue shifts predicted residue probabilities at structurally/sequentially related positions more than random features do.",
        "method": "InterPLM Fig 7 protocol (ESM-2 MLM): clamp a feature on a single unmasked residue, measure the change in masked-token probability at related positions across clamp magnitudes; compare structural, concept-specific, and random features. Optional: RITA generation steering (Villegas Garcia & Ansuini 2025).",
        "rationale": "Steering is the strongest causal-role demonstration in the PLM-SAE literature and yields an intuitive, figure-friendly result.",
        "robustness": "Multiple target features per concept; multiple host sequences; SAE seeds {42,43,44}; random-feature and amino-acid-specific-feature controls; report slope of probability vs clamp magnitude with CIs across sequences.",
        "outputs": "Steering curves (Δp vs clamp magnitude) for structural vs random features; a worked example on a real motif.",
        "feasibility": "Cheap ESM-2 forward passes; <1–2 h. HIGH.",
        "risks": "Effect may be subtle for diffuse features; mitigated by selecting high-F1 features identified in E0.",
    },
    {
        "id": "E4",
        "title": "Null calibration of L_struct  [BUILT: experiment_null_calibration.py]",
        "question": "What is L_struct under structure-destroying nulls, so the effect size can be read against a principled floor?",
        "hypothesis": "L_struct on a degree-matched / globally-shuffled contact graph is ~0; the real-graph effect is many-fold larger.",
        "method": "Recompute the exact L_struct kernel (observed minus within-protein activation shuffle) but replace the real Cα graph with (i) a degree-matched random graph (same degree, same separation filter, random partners within protein) and (ii) a global-shuffle graph (partners anywhere).",
        "rationale": "Answers reviewer U2Yp ('how big is the effect?') by providing the missing lower bound. Pilot: real mean +0.054 vs global-shuffle +0.0007 (~74x), confirming the metric tracks genuine 3D contacts.",
        "robustness": "Run at all depths and both models; >= 3 null seeds; report full per-feature distributions and the real-to-null ratio; combine with C1 randomized-weights floor for a second independent null.",
        "outputs": "Real vs null L_struct histograms; per-model real-to-null ratio vs depth; calibrated restatement of the headline d.",
        "feasibility": "Reuses cached activations + adjacency; ~5 min/layer. HIGH.",
        "risks": "Degree-matched null can go mildly negative (a real, interpretable contrast); the global-shuffle null is the clean ~0 floor and is reported as primary.",
    },
    {
        "id": "E5",
        "title": "Automated feature interpretation with an LLM  [OPTIONAL]",
        "question": "Can features be auto-described and validated predictively, scaling interpretation beyond curated concepts?",
        "hypothesis": "LLM-generated feature descriptions predict held-out maximally-activating proteins above chance, more so for ESM-2 than RITA.",
        "method": "InterPLM auto-interp: summarise a feature's top-activating contexts with an LLM, then test whether the description predicts activation on held-out proteins.",
        "rationale": "Captures concepts absent from Swiss-Prot/SCOPe and provides qualitative interpretability at scale.",
        "robustness": "Held-out predictive validation; multiple features per concept class; blinded human spot-check on a sample (Adams-style ~80% interpretable benchmark).",
        "outputs": "Auto-interp description table; predictive-accuracy distribution; ESM-2 vs RITA comparison.",
        "feasibility": "API-bound, not compute-bound; needs an API key and modest budget. MEDIUM (deferred for v1).",
        "risks": "LLM hallucination; mitigated by quantitative held-out validation rather than trusting descriptions.",
    },
    {
        "id": "E6",
        "title": "Qualitative structural visualisation of top features  [PLANNED]",
        "question": "What do the strongest structural-locality / concept features look like on real protein structures?",
        "hypothesis": "Top features localise to coherent structural elements (β-sheet pairings, buried cores, active sites).",
        "method": "Render top-F1 / top-L_struct features as activation overlays on PDB structures (py3Dmol/PyMOL) for representative proteins.",
        "rationale": "Reviewer 6P8E's request; makes the biology legible and complements quantitative results with intuition.",
        "robustness": "Show features that are stable across SAE seeds; include both a specific (high-precision) and a general (high-recall) example per concept to illustrate the precision/recall trade-off.",
        "outputs": "Figure panels of features on structures with concept labels.",
        "feasibility": "Trivial compute; minutes. HIGH.",
        "risks": "Cherry-picking; mitigated by selecting by pre-registered ranking criteria, not by eye.",
    },
    {
        "id": "E7",
        "title": "Model-panel expansion  [OPTIONAL / heavier]",
        "question": "Does the bidirectional-vs-causal dissociation hold across MULTIPLE models per family, averaging out per-model training noise?",
        "hypothesis": "With >=2 bidirectional and >=2 residue-tokenised causal PLMs, a family-level mixed-effects model shows ESM-family > causal-family on L_struct and concept-F1.",
        "method": "Add clean residue-tokenised comparators; run the full pipeline (embeddings, SAE, E0/E1/E4) per model; fit a mixed-effects model with family as fixed effect and model as random effect.",
        "rationale": "Both reviewers note the single-causal-comparator weakness; a panel converts a two-model contrast into a family-level claim.",
        "robustness": "Each new model inherits the full Section-2 battery (seeds, k, splits, sweeps); report between-model variance explicitly.",
        "outputs": "Per-model L_struct/concept-F1; family-level mixed-effects estimate with CI.",
        "feasibility": "~half-day per new model (full pipeline). MEDIUM; recommended for v2.",
        "risks": "Tokenisation confounds (avoid BPE models for residue-pair metrics, per the paper's own §4.2 finding).",
    },
]

for e in experiments:
    heading(f"{e['id']}. {e['title']}", 2)
    kv("Question", e["question"])
    kv("Hypothesis", e["hypothesis"])
    kv("Method (source)", e["method"])
    kv("Why it is reasonable and scientific", e["rationale"])
    kv("Seed repeats & robustness", e["robustness"])
    kv("Expected outputs", e["outputs"])
    kv("Feasibility (M5 Max 128 GB)", e["feasibility"])
    kv("Risks & mitigations", e["risks"])
    doc.add_paragraph()

# =====================================================================
#  4. FEASIBILITY SUMMARY
# =====================================================================
heading("4. Feasibility summary (MacBook M5 Max, 128 GB)", 1)
table(
    ["Exp", "Compute profile", "Est. wall-clock", "Status", "Risk"],
    [
        ["E0", "CPU on cached activations", "~2.5 h (18 layers)", "BUILT + smoke-tested", "Low"],
        ["C1", "Embedding extraction + SAE training", "~1–2 h/layer", "Planned", "Med"],
        ["E1", "Re-extract embeddings + probe fits", "<1–2 h/model", "BUILT + smoke-tested", "Low"],
        ["E2", "Hooked forward passes (CPU)", "Overnight", "Extends clamping", "Med"],
        ["E3", "ESM-2 masked-prob forward passes", "<1–2 h", "Extends clamping", "Low"],
        ["E4", "Cached activations + adjacency", "~5 min/layer", "BUILT + smoke-tested", "Low"],
        ["E5", "LLM API", "API-bound", "Optional (needs key)", "Med"],
        ["E6", "Structure rendering", "Minutes", "Planned", "Low"],
        ["E7", "Full pipeline per new model", "~half-day each", "Optional (v2)", "Med"],
    ])

# =====================================================================
#  5. PHASED EXECUTION (robust-first)
# =====================================================================
heading("5. Phased execution timeline (robustness built into each phase)", 1)
para("Robustness is not a final phase; each phase below runs its experiments through the relevant Section-2 checks before the phase is declared done.")
bullets([
    ("Phase 1 — Validity, necessity, calibration (built)", "E0 concept-F1, E1 SAE-vs-raw probe, E4 null calibration, plus C1 randomized-weights control. Run all at SAE seeds {42,43,44} and confirm on the 150-protein val split. Deliverable: L_struct grounded in curated biology, shown necessary vs a linear baseline, and calibrated against a near-zero floor."),
    ("Phase 2 — Causality (priority)", "E2 causal-feature identification + causal-H1 and E3 steering, both with paired bootstrap CIs and random-feature controls. Deliverable: the dissociation holds on the causally load-bearing subset."),
    ("Phase 3 — Breadth and polish", "E6 qualitative structures, E5 auto-interp (if API available), E7 model-panel expansion with a family-level mixed-effects analysis."),
])

# =====================================================================
#  6. ACCEPTANCE / WHAT MAKES A CLAIM PUBLISHABLE
# =====================================================================
heading("6. Acceptance criteria for a reported claim", 1)
bullets([
    "Direction of the effect is preserved across SAE seeds {42,43,44} with small cross-seed SD.",
    "95% protein-level cluster-bootstrap CI excludes 0 at the matched depths cited.",
    "Effect is preserved on the 150-protein held-out validation set.",
    "Effect collapses to floor under the randomized-weights control (C1) and shuffled-graph null (E4).",
    "Many-feature analyses report Benjamini–Hochberg FDR and use val→test selection to avoid leakage.",
    "Convergence: the bidirectional-vs-causal dissociation is independently visible on >=2 of {L_struct, concept-F1, probe AUROC, causal effect}.",
])

# =====================================================================
#  7. REFERENCES
# =====================================================================
heading("7. Key references", 1)
bullets([
    "Simon & Zou. InterPLM: discovering interpretable features in protein language models via sparse autoencoders. Nature Methods, 2025.",
    "Adams et al. From Mechanistic Interpretability to Mechanistic Biology: Training, Evaluating, and Interpreting Sparse Autoencoders on Protein Language Models. ICML, 2025.",
    "Villegas Garcia & Ansuini. Interpreting and Steering Protein Language Models through Sparse Autoencoders. arXiv:2502.09135, 2025.",
    "Gao et al. Scaling and evaluating sparse autoencoders. arXiv:2406.04093, 2024.",
    "Bricken et al. Towards Monosemanticity. Anthropic, 2023.",
    "Lin et al. (ESM-2), Science 2023; Hesslow et al. (RITA), 2022; Elnaggar et al. (ProtT5), 2021.",
])

out = "EXPERIMENT_PLAN.docx"
doc.save(out)
print(f"Wrote {out}")

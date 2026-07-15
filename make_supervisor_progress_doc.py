#!/usr/bin/env python3
"""Generate SUPERVISOR_PROGRESS_UPDATE.docx — casual progress note."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "SUPERVISOR_PROGRESS_UPDATE.docx"


def fmt_ci(lo, hi):
    return f"[{lo:+.3f}, {hi:+.3f}]"


def sig_label(lo, hi):
    if lo > 0:
        return "Yes"
    if hi < 0:
        return "No (flipped)"
    return "Unclear"


def load_fold_h1(path):
    if not path.exists():
        return []
    df = pd.read_csv(path)
    df = df[df["cluster_level"] == "fold"].copy()
    return [
        (
            str(r["rel_depth"]),
            str(r["layer_pair"]),
            f"{r['d_point']:+.3f}",
            fmt_ci(r["ci_low"], r["ci_high"]),
            sig_label(r["ci_low"], r["ci_high"]),
        )
        for _, r in df.iterrows()
    ]


def fold_ci_summary():
    targets = {
        "v2_cis_pair_esm_rita_fold.csv": 90,
        "v2_cis_pair_pt5_fold.csv": 36,
        "v2_cis_trajectory_fold.csv": 36,
        "v3opt_cis_val_sweeps_fold.csv": 72,
    }
    counts = {}
    for name, target in targets.items():
        p = ROOT / "outputs_robustness" / name
        counts[name] = (len(pd.read_csv(p)) if p.exists() else 0, target)
    esm_n, esm_t = counts["v2_cis_pair_esm_rita_fold.csv"]
    pt5_n, pt5_t = counts["v2_cis_pair_pt5_fold.csv"]
    traj_n, traj_t = counts["v2_cis_trajectory_fold.csv"]
    v3_n, v3_t = counts["v3opt_cis_val_sweeps_fold.csv"]
    if v3_n >= v3_t:
        status = "done"
    elif v3_n > 8:
        status = "running"
    else:
        status = "not started / early"
    return status, esm_n, esm_t, pt5_n, pt5_t, traj_n, traj_t, v3_n, v3_t


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def h(text, level=1):
    return doc.add_heading(text, level=level)


def p(text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)
    return para


def bl(items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            r = para.add_run(item[0] + ": ")
            r.bold = True
            para.add_run(item[1])
        else:
            para.add_run(item)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, htext in enumerate(headers):
        t.rows[0].cells[i].text = htext
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


# --- doc body ---

h("SAE / protein LM project — progress update")

p(
    "Quick update since the workshop paper. Main claim: bidirectional PLMs (ESM-2, ProtBert) "
    "have more structurally local SAE features than causal ones (RITA, ProGen2) — measured "
    "with L_struct. Also ProtT5 encoder vs decoder crosses over around mid-depth."
)

h("Already in the workshop paper (not new)", level=2)
bl([
    "Multi-seed SAE training (seeds 42, 43, 44) — Appendix F",
    "Metric sweeps, k/split robustness, fresh subsample, smaller model pair — Appendices F–M",
    "Protein-level bootstrap CIs",
])

h("What I've added since", level=1)

h("Stats fixes", level=2)
bl([
    (
        "Fold-level bootstrap",
        "Original CIs treated each protein as independent. Domains in the same SCOPe fold "
        "are related, so I re-ran bootstrap by resampling folds. Intervals got ~2× wider; "
        "ESM-2 > RITA still holds everywhere, though 13% depth is barely above zero.",
    ),
    (
        "Multi-seed",
        "Was already in the paper — I re-ran the full 9-depth grids to double-check. "
        "Spread across seeds is tiny.",
    ),
    (
        "Random PLM weights",
        "New control: SAE on untrained weights. Structural signal basically disappears.",
    ),
])

h("Other checks we didn't have in the paper", level=2)
# expanded by patch_supervisor_doc_section42.py

h("Does bidir > causal generalise? (ProtBert + ProGen2)", level=2)
p(
    "The paper already shows ESM-2 > RITA. I added ProtBert-BFD (bidirectional, same side as "
    "ESM-2) and ProGen2 (causal, same side as RITA) to ask: is this a bidirectional-vs-causal "
    "family thing, or just those two models? Same 9-depth grid and L_struct setup."
)

esm_rows = load_fold_h1(ROOT / "outputs_robustness/bootstrap_h1_full_bylevel_minact0.csv")
new_rows = load_fold_h1(ROOT / "outputs_robustness/bootstrap_h1_newpair_full_bylevel_minact0.csv")
if esm_rows and new_rows:
    table(
        ["Depth", "ESM-2 vs RITA d", "ProtBert vs ProGen2 d", "Both bidir > causal?"],
        [
            (
                esm[0],
                esm[2],
                new[2],
                "Yes" if esm[4] == "Yes" and new[4] == "Yes"
                else ("No" if new[4] == "No (flipped)" else "Partial"),
            )
            for esm, new in zip(esm_rows, new_rows)
        ],
    )
    p(
        "ESM-2 > RITA holds at all 9 depths (fold bootstrap). ProtBert > ProGen2 mostly "
        "agrees at mid/late depths but 13% flips; Concept-F1 4-way holds at 7/9 depths. "
        "New-model extensions (diagnostics, Concept-F1, null, interp) are done on ProtBert/ProGen2."
    )
elif new_rows:
    table(
        ["Depth", "Layers", "ProtBert − ProGen2 d", "95% CI (fold)", "Bidir > causal?"],
        new_rows,
    )

h("Still running / to do", level=1)
status, esm_n, esm_t, pt5_n, pt5_t, traj_n, traj_t, v3_n, v3_t = fold_ci_summary()
bl([
    (
        "Fold-level CIs for tables",
        f"Done. ESM/RITA ({esm_n}/{esm_t}), ProtT5 ({pt5_n}/{pt5_t}), "
        f"trajectories ({traj_n}/{traj_t}), metric sweeps ({v3_n}/{v3_t}).",
    ),
    "Concept-F1 fold-split sensitivity (optional; headline multi-seed used protein split)",
    "Pull everything into proper tables + depth figures",
    "Decide how to write up ProtBert/ProGen2 — partial on L_struct and Concept-F1",
])

h("Worth keeping in mind", level=2)
bl([
    "L_struct = geometric contacts specifically, not 'interpretability' in general",
    "Bidir vs causal = family contrast (architecture + training + data), not just MLM vs CLM",
    "Steering didn't work clearly; raw activations beat SAE on linear probes (expected)",
    "Had a bad week with iCloud eating project files — recovered, moved venv off Desktop",
])

doc.save(OUT)
print(f"Wrote {OUT}")

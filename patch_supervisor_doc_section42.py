#!/usr/bin/env python3
"""Insert §4.2 body into SUPERVISOR_PROGRESS_UPDATE.docx."""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
DOC_PATH = ROOT / "SUPERVISOR_PROGRESS_UPDATE.docx"

BLOCKS = [
    ("normal", (
        "Workshop paper leaned on one metric (L_struct). Since then I've checked the "
        "ESM-2 vs RITA gap several other ways, so the result doesn't rest on a single "
        "homemade number. Each check below: why I ran it, how it works, what came out."
    )),
    ("normal", (
        "Extension checks below use SAE seed 42 unless noted. Concept-F1 headline cells "
        "now also have SAE seeds 43/44 (see multi-seed table). L_struct multi-seed was "
        "already in the workshop paper."
    )),
    ("h3", "Annotation matching (Concept-F1)"),
    ("bullet", ("Why", (
        "L_struct is something we invented. This checks whether SAE features line up with "
        "labels biologists already use (SCOPe fold/family, helix/strand, buried vs exposed) — "
        "the standard InterPLM test."
    ))),
    ("bullet", ("How", (
        "Treat each feature as a yes/no detector with a threshold. On one set of proteins, "
        "find the best feature+threshold for each label; then score its F1 on a separate, "
        "held-out set of proteins (split by fold so close relatives can't leak). High F1 "
        "means the feature really tracks the label across proteins, not memorised ones."
    ))),
    ("bullet", ("Threshold", (
        "Activations max-normalised to [0,1]; best threshold per concept picked from "
        "{0, 0.15, 0.5, 0.6, 0.8} on the validation set, then fixed for test. "
        "Caveat: this is a coarse selection grid — I haven't run a dedicated sensitivity "
        "sweep over it (L_struct did get threshold sweeps in the paper; Concept-F1 hasn't)."
    ))),
    ("bullet", ("Result", (
        "ESM-2 peaks ~0.73 (L16), RITA ~0.64 (L18). Same direction as L_struct. "
        "ProtBert/ProGen2 Concept-F1 now run too (9 depths each); 4-way bidir-vs-causal "
        "on mean F1 holds at 7/9 depths (fails at 13% and 100%) — still weaker story than "
        "L_struct at every depth. "
        "Headline multi-seed (protein split, 80 concepts): trained ESM-2 L16 "
        "0.714±0.009 (SAE seeds 42/43/44) vs random weights 0.30±0.13 (PLM weight seeds 0/1/2)."
    ))),
    ("h3", "Linear probes"),
    ("bullet", ("Why", (
        "Sanity check that the ESM-2 vs RITA gap isn't an artefact of our metric — does it "
        "show up on a bog-standard 'predict the structure label' task?"
    ))),
    ("bullet", ("How", (
        "Train simple classifiers (helix / strand / buried-vs-exposed) on raw PLM vectors, "
        "and separately on SAE features, then compare accuracy."
    ))),
    ("bullet", ("Result", (
        "Raw PLM beats SAE for decoding — expected, SAEs aren't built for prediction, so this "
        "is a different job, not an SAE failure. The useful bit: ESM-2 raw beats RITA raw at "
        "every depth, same direction, without using L_struct."
    ))),
    ("h3", "Null baseline"),
    ("bullet", ("Why", (
        "An effect size means nothing without knowing what 'zero' looks like — how far above "
        "chance is L_struct really?"
    ))),
    ("bullet", ("How", (
        "Recompute L_struct against shuffled contact graphs (random structure) and take the "
        "ratio of real to shuffled."
    ))),
    ("bullet", ("Result", (
        "ESM-2 sits ~150–1350× above the null floor; RITA ~10–22×. Both beat chance, ESM-2 by "
        "far more."
    ))),
    ("h3", "Do the same features score highly on everything?"),
    ("bullet", ("Why", (
        "If L_struct, Concept-F1 and burial/exposure all rank the same features, they're "
        "redundant. If not, they're independent evidence and L_struct shouldn't be sold as "
        "general 'interpretability'."
    ))),
    ("bullet", ("How", (
        "For each feature (same dictionary throughout), score it under L_struct, under "
        "Concept-F1, and by how well it tracks residue exposure (SASA) and neighbour count, "
        "then rank-correlate those scores."
    ))),
    ("bullet", ("Result", (
        "Mostly unrelated (ρ ~ 0.01–0.28). A feature can be structurally local without "
        "matching a SCOPe label, and vice versa — the checks really are independent."
    ))),
    ("h3", "SAE health + faithfulness"),
    ("bullet", ("Why", (
        "Make sure the SAEs aren't broken (dead/duplicate features) and that they actually "
        "preserve what the model does."
    ))),
    ("bullet", ("How", (
        "Check dictionary redundancy and reconstruction quality; then swap SAE reconstructions "
        "back into the model and see how much behaviour is kept."
    ))),
    ("bullet", ("Result", (
        "Dictionaries fine. RITA reconstructs almost everything (high variance explained). "
        "Putting SAE back into ESM-2: easy at shallow layers, harder deep."
    ))),
    ("h3", "Causal tests (early days)"),
    ("bullet", ("Why", (
        "Everything above is correlational. This asks whether structural features actually do "
        "anything — if you remove or push them, does structure-related behaviour change?"
    ))),
    ("bullet", ("How", (
        "Ablation: zero out top structural features vs random ones and watch contact readout. "
        "Steering: clamp a feature up/down and see if contact statistics shift."
    ))),
    ("bullet", ("Result", (
        "Ablation at ESM-2 L16 hurts a bit more than random (p=0.02) but tiny. Steering: right "
        "direction, not significant (p~0.1). Wouldn't lean on either yet."
    ))),
    ("h3", "Random PLM weights (negative control)"),
    ("bullet", ("Why", (
        "Confirms the signal comes from the trained model, not from the SAE machinery dressing "
        "up noise."
    ))),
    ("bullet", ("How", (
        "Train the same SAE on an untrained (randomly initialised) PLM and rerun the checks."
    ))),
    ("bullet", ("Result", (
        "Structural alignment basically vanishes; only coarse amino-acid composition survives. "
        "So you need a real trained model."
    ))),
    ("normal", (
        "Bottom line: annotation matching and probes both back up ESM-2 > RITA without "
        "L_struct. Causal/steering is weak so far."
    )),
]


def make_paragraph(doc, text="", style=None):
    para = doc.add_paragraph(text, style=style)
    el = para._element
    doc.element.body.remove(el)
    return el


def add_bullet(doc, content):
    para = doc.add_paragraph(style="List Bullet")
    if isinstance(content, tuple):
        label, body = content
        if label:
            r = para.add_run(f"{label}: ")
            r.bold = True
        para.add_run(body)
    else:
        para.add_run(content)
    el = para._element
    doc.element.body.remove(el)
    return el


def find_para_index(doc, prefix):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i
    raise ValueError(f"not found: {prefix!r}")


def patch():
    doc = Document(DOC_PATH)
    i42 = find_para_index(doc, "Other checks we didn't have in the paper")
    i43 = find_para_index(doc, "Does bidir > causal generalise")

    body = doc.element.body
    children = list(body)
    el42 = doc.paragraphs[i42]._element
    el43 = doc.paragraphs[i43]._element
    idx42 = children.index(el42)
    idx43 = children.index(el43)

    for el in children[idx42 + 1 : idx43]:
        body.remove(el)

    new_elements = []
    for kind, content in BLOCKS:
        if kind == "normal":
            el = make_paragraph(doc, content, style="Normal")
        elif kind == "h3":
            el = make_paragraph(doc, content, style="Heading 3")
        elif kind == "bullet":
            el = add_bullet(doc, content)
        new_elements.append(el)

    anchor = el42
    for el in new_elements:
        anchor.addnext(el)
        anchor = el

    doc.save(DOC_PATH)
    print(f"Patched §4.2 in {DOC_PATH} ({len(new_elements)} paragraphs)")


if __name__ == "__main__":
    patch()
